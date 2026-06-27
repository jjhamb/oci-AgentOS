#!/usr/bin/env python3
"""Generate EC Minutes DOCX — v3, one page, proper spacing."""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup — A4, moderate margins ──
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ── Helpers ──
def set_run_font(run, font_name='Times New Roman', font_size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_para(text, align='left', font_size=12, bold=False, space_after=2, space_before=0, indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.15
    if indent is not None:
        pf.left_indent = Cm(indent)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', font_size, bold)
    return p

def add_run(p, text, font_size=12, bold=False):
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', font_size, bold)
    return run

# ══════════════════════════════════════════════
# LETTERHEAD — positioned higher (top of page)
# ══════════════════════════════════════════════
header_img = '/home/jayant/Desktop/Hermes-AgentOS-Dashboard/tabs/mezzaria-style/image1.png'
footer_img = '/home/jayant/Desktop/Hermes-AgentOS-Dashboard/tabs/mezzaria-style/image2.png'

if os.path.exists(header_img):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    run = p.add_run()
    run.add_picture(header_img, width=Inches(5.5))
else:
    add_para('MEZZARIA', align='center', font_size=14, bold=True, space_after=1)
    add_para("FLAT BUYERS' WELFARE ASSOCIATION", align='center', font_size=12, bold=True, space_after=1)

# Contact line — right aligned, tight
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(1)
add_run(p, 'e-Mail: mezzariafbwa@gmail.com', font_size=9)
p.add_run('    ')
add_run(p, 'Web Site: www.mezzaria.org', font_size=9)

# ══════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════
add_para('MINUTES OF MEETING', align='center', font_size=16, bold=True, space_after=1)
add_para('Executive Committee (EC)', align='center', font_size=12, bold=True, space_after=0)
add_para('Mezzaria Flat Buyers Welfare Association', align='center', font_size=11, space_after=2)

# Meeting details — compact
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(3)
add_run(p, 'Date: 13th June 2026', font_size=11)
p.add_run('  |  ')
add_run(p, 'Time: 1500 hrs', font_size=11)
p.add_run('  |  ')
add_run(p, 'Venue: Mezzaria Club House', font_size=11)

# ══════════════════════════════════════════════
# ATTENDEES
# ══════════════════════════════════════════════
add_para('Attendees:', font_size=11, bold=True, space_after=2, space_before=4)

attendees = [
    '1.  Rahoul Mundra',
    '2.  Suneal Singhal',
    '3.  Meenakshi Jhamb',
    '4.  Jayant Jhamb',
    '5.  Narendra Kumar',
    '6.  Gp Capt H.K. Srivastava',
    '7.  Shwetank Garg (online — Zoom)',
    '8.  Dinesh Sharma (online — Zoom, briefly)',
]
for a in attendees:
    add_para(a, font_size=11, space_after=0, indent=0.5)

# ══════════════════════════════════════════════
# DISCUSSION
# ══════════════════════════════════════════════
add_para('', font_size=11, space_before=4)
add_para('Points discussed:', font_size=11, bold=True, space_after=3)

points = [
    ('1.', 'Acceptance of resignation by Sinha sahab.',
     'The Executive Committee accepted the resignation of Gurudeo Sinha from the post of President. Sinha sahab cited reasons of not keeping in good health, which is impairing his ability to actively participate in the affairs of the Association. The Committee placed on record its appreciation for his contributions during his tenure.'),
    ('2.', 'Proposal of making Jayant Jhamb President.',
     'Shwetank Garg proposed the name of Capt. Jayant Jhamb for the post of President. The proposal was seconded by Suneal Singhal. All present members approved the appointment, and Capt. Jhamb graciously accepted the responsibility.'),
    ('3.', 'Proposal to choose another person as treasurer.',
     'Meenakshi Jhamb expressed her inability to continue as Treasurer. The Committee discussed the need to identify and appoint a new Treasurer.'),
    ('4.', 'Association to pursue delinking of electricity meter from maintenance.',
     'The Committee discussed and agreed that the Association should actively pursue the delinking of electricity meters from maintenance charges.'),
    ('5.', 'KYC completion and payments.',
     'It was informed that KYC documentation will be completed within a day or two. Accordingly, payments will be carried out as requisitioned once KYC is in order.'),
]

for num, title, detail in points:
    # Number + title in bold
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(1)
    pf.space_before = Pt(2)
    pf.line_spacing = 1.15
    pf.left_indent = Cm(0.5)
    run = p.add_run(f'{num}  ')
    set_run_font(run, 'Times New Roman', 11, True)
    run2 = p.add_run(title)
    set_run_font(run2, 'Times New Roman', 11, True)
    # Detail in regular
    if detail:
        p2 = doc.add_paragraph()
        pf2 = p2.paragraph_format
        pf2.space_after = Pt(1)
        pf2.line_spacing = 1.15
        pf2.left_indent = Cm(1.2)
        run3 = p2.add_run(detail)
        set_run_font(run3, 'Times New Roman', 11, False)

# ══════════════════════════════════════════════
# SIGNATORY BLOCK
# ══════════════════════════════════════════════
add_para('', font_size=11, space_before=6)

table = doc.add_table(rows=3, cols=2)

# Add visible borders to the table
from docx.oxml import OxmlElement
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
borders = OxmlElement('w:tblBorders')
for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
    element = OxmlElement(f'w:{edge}')
    element.set(qn('w:val'), 'single')
    element.set(qn('w:sz'), '4')
    element.set(qn('w:space'), '0')
    element.set(qn('w:color'), '000000')
    borders.append(element)
tblPr.append(borders)

# Reduce cell spacing
for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for edge in ('top', 'left', 'bottom', 'right'):
            mar = OxmlElement(f'w:{edge}')
            mar.set(qn('w:w'), '40')
            mar.set(qn('w:type'), 'dxa')
            tcMar.append(mar)
        tcPr.append(tcMar)

# Row 2: Signatories
cells = table.rows[1].cells
p_left = cells[0].paragraphs[0]
p_left.paragraph_format.space_after = Pt(1)
add_run(p_left, 'Capt. Jayant Jhamb', font_size=11, bold=True)
p_left2 = cells[0].add_paragraph()
p_left2.paragraph_format.space_after = Pt(1)
add_run(p_left2, 'President', font_size=11)
p_left3 = cells[0].add_paragraph()
p_left3.paragraph_format.space_after = Pt(0)
add_run(p_left3, 'Mobile: 8826055550', font_size=10)

p_right = cells[1].paragraphs[0]
p_right.paragraph_format.space_after = Pt(1)
add_run(p_right, 'Suneal Kumar Singhal', font_size=11, bold=True)
p_right2 = cells[1].add_paragraph()
p_right2.paragraph_format.space_after = Pt(1)
add_run(p_right2, 'Secretary', font_size=11)
p_right3 = cells[1].add_paragraph()
p_right3.paragraph_format.space_after = Pt(0)
add_run(p_right3, 'Mobile: 9811428209', font_size=10)

# Row 3: Association line
cells = table.rows[2].cells
p_assoc = cells[0].paragraphs[0]
p_assoc.paragraph_format.space_before = Pt(2)
add_run(p_assoc, 'On behalf of Mezzaria Flat Buyers Welfare Association', font_size=10)
cells[1].text = ''

# ══════════════════════════════════════════════
# FOOTER IMAGE — positioned lower
# ══════════════════════════════════════════════
if os.path.exists(footer_img):
    add_para('', font_size=11, space_before=4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(footer_img, width=Inches(5.5))

# ══════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════
output = '/home/jayant/Desktop/Hermes-AgentOS-Dashboard/tabs/mezzaria-style/EC_Minutes_13Jun2026.docx'
doc.save(output)
print(f'Saved: {output}')
