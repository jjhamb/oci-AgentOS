#!/usr/bin/env python3
"""Generate EC Minutes DOCX in Mezzaria style."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Helper functions ──
def set_run_font(run, font_name='Times New Roman', font_size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_paragraph(text, align='left', font_size=12, bold=False, font='Times New Roman', space_after=6, space_before=0, first_line_indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.5
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, font, font_size, bold)
    return p

def add_run_to_paragraph(p, text, font_size=12, bold=False, font='Times New Roman'):
    run = p.add_run(text)
    set_run_font(run, font, font_size, bold)
    return run

# ══════════════════════════════════════════════
# HEADER / LETTERHEAD
# ══════════════════════════════════════════════

# Try to add the header image
header_img_path = '/home/jayant/Desktop/Hermes-AgentOS-Dashboard/tabs/mezzaria-style/image1.png'
if os.path.exists(header_img_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(header_img_path, width=Inches(5.0))
else:
    # Fallback: text-based letterhead
    add_paragraph('MEZZARIA', align='center', font_size=16, bold=True, space_after=2)
    add_paragraph("FLAT BUYERS' WELFARE ASSOCIATION", align='center', font_size=14, bold=True, space_after=2)
    add_paragraph('E-611, First Floor, Greater Kailash Part - II, New Delhi – 110048', align='center', font_size=10, space_after=2)

# Contact line (right-aligned email + website)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run_to_paragraph(p, 'e-Mail: mezzariafbwa@gmail.com', font_size=10)
p.add_run('    ')
add_run_to_paragraph(p, 'Web Site: www.mezzaria.org', font_size=10)

# Separator line
add_paragraph('─' * 80, align='center', font_size=8, space_after=4)

# ══════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════
add_paragraph('MINUTES OF MEETING', align='center', font_size=16, bold=True, space_after=4)
add_paragraph('Executive Committee (EC) Meeting', align='center', font_size=12, bold=True, space_after=2)
add_paragraph('Held on 13th June 2025', align='center', font_size=12, space_after=8)

# ══════════════════════════════════════════════
# ATTENDEES
# ══════════════════════════════════════════════
add_paragraph('Attendees:', font_size=12, bold=True, space_after=4)

attendees = [
    '1.  Rahoul Mundra',
    '2.  Suneal Singhal',
    '3.  Meenakshi Jhamb',
    '4.  Jayant Jhamb',
    '5.  Narendra Kumar (on phone)',
    '6.  Gp Capt HK Srivastava',
]

for a in attendees:
    add_paragraph(a, font_size=12, space_after=2, first_line_indent=0.5)

add_paragraph('Mom awaited.', font_size=12, space_after=8)

# ══════════════════════════════════════════════
# POINTS DISCUSSED
# ══════════════════════════════════════════════
add_paragraph('Points discussed on 13th June 2025:', font_size=12, bold=True, space_after=4)

points = [
    ('1.', 'Acceptance of resignation by Sinha sahab.'),
    ('2.', 'Proposal of making Jayant Jhamb President.'),
    ('3.', 'Proposal to choose another person as treasurer as I can\'t be anymore.'),
    ('4.', 'Association to pursue delinking of electricity meter from maintenance.'),
    ('5.', 'Meanwhile KYC will be done in a day or two, so payments will be carried out as requisitioned.'),
]

for num, text in points:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(2)
    pf.line_spacing = 1.5
    pf.left_indent = Cm(0.5)
    run = p.add_run(f'{num}  ')
    set_run_font(run, 'Times New Roman', 12, True)
    run2 = p.add_run(text)
    set_run_font(run2, 'Times New Roman', 12, False)

# ══════════════════════════════════════════════
# FOOTER IMAGE
# ══════════════════════════════════════════════
footer_img_path = '/home/jayant/Desktop/Hermes-AgentOS-Dashboard/tabs/mezzaria-style/image2.png'
if os.path.exists(footer_img_path):
    # Add some space before footer
    add_paragraph('', font_size=12, space_after=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(footer_img_path, width=Inches(5.0))

# ══════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════
output_path = '/home/jayant/Desktop/Hermes-AgentOS-Dashboard/tabs/mezzaria-style/EC_Minutes_13Jun2025.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
