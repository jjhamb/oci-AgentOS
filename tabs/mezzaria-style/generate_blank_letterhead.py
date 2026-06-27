#!/usr/bin/env python3
"""Generate blank Mezzaria letterhead template DOCX."""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup — A4 ──
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

def set_run_font(run, font_name='Times New Roman', font_size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_para(text, align='left', font_size=12, bold=False, space_after=2, space_before=0, indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.15
    if indent is not None:
        pf.left_indent = Cm(indent)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', font_size, bold)
    return p

def add_run(p, text, font_size=12, bold=False):
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', font_size, bold)
    return run

# ══════════════════════════════════════════════
# LETTERHEAD
# ══════════════════════════════════════════════
header_img = '/home/jayant/Desktop/Hermes-AgentOS-Dashboard/tabs/mezzaria-style/image1.png'
footer_img = '/home/jayant/Desktop/Hermes-AgentOS-Dashboard/tabs/mezzaria-style/image2.png'

if os.path.exists(header_img):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    run = p.add_run()
    run.add_picture(header_img, width=Inches(5.5))

# ══════════════════════════════════════════════
# BLANK BODY — single empty paragraph for content
# ══════════════════════════════════════════════
add_para('', font_size=11, space_before=4)

# ══════════════════════════════════════════════
# FOOTER IMAGE
# ══════════════════════════════════════════════
if os.path.exists(footer_img):
    add_para('', font_size=11, space_before=4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(footer_img, width=Inches(5.5))

output = '/home/jayant/Desktop/Hermes-AgentOS-Dashboard/tabs/mezzaria-style/Mezzaria_Letterhead_Blank.docx'
doc.save(output)
print(f'Saved: {output}')
