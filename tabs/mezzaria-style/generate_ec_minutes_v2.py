#!/usr/bin/env python3
"""Generate EC Minutes DOCX in Mezzaria style — v2, properly formatted."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Helpers ──
def set_run_font(run, font_name='Times New Roman', font_size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_para(text, align='left', font_size=12, bold=False, space_after=4, space_before=0, indent=None, font_name='Times New Roman'):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.5
    if indent is not None:
        pf.left_indent = Cm(indent)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_font(run, font_name, font_size, bold)
    return p

def add_run(p, text, font_size=12, bold=False):
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', font_size, bold)
    return run

def add_page_number(paragraph):
    """Add page number field to a paragraph."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    run2._r.append(instrText)
    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar2)
    # Add " of " and total
    run4 = paragraph.add_run(' of ')
    set_run_font(run4, 'Times New Roman', 10)
    run5 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run5._r.append(fldChar3)
    run6 = paragraph.add_run()
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = ' NUMPAGES '
    run6._r.append(instrText2)
    run7 = paragraph.add_run()
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run7._r.append(fldChar4)

# ══════════════════════════════════════════════
# LETTERHEAD — use image at original proportions
# ══════════════════════════════════════════════
header_img = '/home/jayant/Desktop/Hermes-AgentOS-Dashboard/tabs/mezzaria-style/image1.png'
footer_img = '/home/jayant/Desktop/Hermes-AgentOS-Dashboard/tabs/mezzaria-style/image2.png'

if os.path.exists(header_img):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # Original is 697x103 px — use width that fits page (~6.5 inches) keeping aspect
    run.add_picture(header_img, width=Inches(6.0))
    pf = p.paragraph_format
    pf.space_after = Pt(2)
else:
    add_para('MEZZARIA', align='center', font_size=16, bold=True, space_after=2)
    add_para("FLAT BUYERS' WELFARE ASSOCIATION", align='center', font_size=14, bold=True, space_after=2)

# Contact line — right aligned
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p, 'e-Mail: mezzariafbwa@gmail.com', font_size=10)
p.add_run('    ')
add_run(p, 'Web Site: www.mezzaria.org', font_size=10)

# Thin separator
add_para('─' * 90, align='center', font_size=6, space_after=2, space_before=2)

# ══════════════════════════════════════════════
# TITLE — bold, centered, large
# ══════════════════════════════════════════════
add_para('MINUTES OF MEETING', align='center', font_size=18, bold=True, space_after=2)
add_para('Executive Committee (EC)', align='center', font_size=13, bold=True, space_after=1)
add_para('Mezzaria Flat Buyers Welfare Association', align='center', font_size=12, space_after=4)

# ══════════════════════════════════════════════
# MEETING DETAILS
# ══════════════════════════════════════════════
add_para('Date: 13th June 2026', font_size=12, space_after=1)
add_para('Time: 1500 hrs', font_size=12, space_after=1)
add_para('Venue: Mezzaria Club House', font_size=12, space_after=6)

# ══════════════════════════════════════════════
# ATTENDEES
# ══════════════════════════════════════════════
add_para('Attendees:', font_size=12, bold=True, space_after=3)

attendees = [
    '1.  Rahoul Mundra',
    '2.  Suneal Singhal',
    '3.  Meenakshi Jhamb',
    '4.  Jayant Jhamb',
    '5.  Narendra Kumar (on phone)',
    '6.  Gp Capt H.K. Srivastava',
]
for a in attendees:
    add_para(a, font_size=12, space_after=1, indent=0.5)

# Leave of absence
add_para('', font_size=12, space_after=1)
add_para('Leave of Absence:', font_size=12, bold=True, space_after=1)
add_para('Mom', font_size=12, space_after=6, indent=0.5)

# ══════════════════════════════════════════════
# DISCUSSION
# ══════════════════════════════════════════════
add_para('The following points were discussed in the meeting:', font_size=12, space_after=4)

# Points with proper formatting
points = [
    ('1.', 'Acceptance of resignation by Sinha sahab.',
     'The Executive Committee accepted the resignation of Gurudeo Sinha from the post of President. Sinha sahab cited reasons of not keeping in good health, which is impairing his ability to actively participate in the affairs of the Association. The Committee placed on record its appreciation for his contributions during his tenure.'),
    ('2.', 'Proposal of making Jayant Jhamb President.',
     'The Committee proposed the name of Capt. Jayant Jhamb for the post of President of the Association. The proposal was discussed and [to be confirmed — was this unanimously approved?].'),
    ('3.', 'Proposal to choose another person as treasurer.',
     'Meenakshi Jhamb expressed her inability to continue as Treasurer. The Committee discussed the need to identify and appoint a new Treasurer. [To be confirmed — was a replacement named or is the matter still open?].'),
    ('4.', 'Association to pursue delinking of electricity meter from maintenance.',
     'The Committee discussed and agreed that the Association should actively pursue the delinking of electricity meters from maintenance charges. [To be confirmed — who will lead this initiative and what is the proposed timeline?].'),
    ('5.', 'KYC completion and payments.',
     'It was informed that KYC documentation will be completed within a day or two. Accordingly, payments will be carried out as requisitioned once KYC is in order.'),
]

for num, title, detail in points:
    # Number + title in bold
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(4)
    pf.line_spacing = 1.5
    pf.left_indent = Cm(0.5)
    run = p.add_run(f'{num}  ')
    set_run_font(run, 'Times New Roman', 12, True)
    run2 = p.add_run(title)
    set_run_font(run2, 'Times New Roman', 12, True)
    # Detail in regular
    if detail:
        p2 = doc.add_paragraph()
        pf2 = p2.paragraph_format
        pf2.space_after = Pt(4)
        pf2.line_spacing = 1.5
        pf2.left_indent = Cm(1.2)
        run3 = p2.add_run(detail)
        set_run_font(run3, 'Times New Roman', 12, False)

# ══════════════════════════════════════════════
# CLARITY NEEDED — notes section
# ══════════════════════════════════════════════
add_para('', font_size=12, space_after=2)
p_note = doc.add_paragraph()
pf_note = p_note.paragraph_format
pf_note.space_before = Pt(6)
pf_note.space_after = Pt(2)
run_note = p_note.add_run('Points requiring clarity / follow-up:')
set_run_font(run_note, 'Times New Roman', 12, True)

clarifications = [
    '• Was the proposal to appoint Jayant Jhamb as President unanimously approved?',
    '• Has a replacement Treasurer been identified, or is the matter still open?',
    '• Who will lead the electricity meter delinking initiative?',
    '• Any specific timeline for the delinking pursuit?',
]
for c in clarifications:
    add_para(c, font_size=11, space_after=1, indent=1.0, font_name='Times New Roman')

# ══════════════════════════════════════════════
# SIGNATORY BLOCK
# ══════════════════════════════════════════════
add_para('', font_size=12, space_after=12)

# Two-column signatory
table = doc.add_table(rows=3, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Row 1: Names
cells = table.rows[0].cells
cells[0].text = ''
cells[1].text = ''
# Row 2: Left signatory
cells = table.rows[1].cells
p_left = cells[0].paragraphs[0]
p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p_left, 'Capt. Jayant Jhamb', font_size=12, bold=True)
p_left2 = cells[0].add_paragraph()
p_left2.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p_left2, 'President', font_size=12)
p_left3 = cells[0].add_paragraph()
p_left3.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p_left3, 'Mobile: 8826055550', font_size=11)

# Right signatory
p_right = cells[1].paragraphs[0]
p_right.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p_right, 'Suneal Kumar Singhal', font_size=12, bold=True)
p_right2 = cells[1].add_paragraph()
p_right2.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p_right2, 'Secretary', font_size=12)
p_right3 = cells[1].add_paragraph()
p_right3.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p_right3, 'Mobile: 9811428209', font_size=11)

# Row 3: Association line
cells = table.rows[2].cells
p_assoc = cells[0].paragraphs[0]
p_assoc.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p_assoc, 'On behalf of Mezzaria Flat Buyers Welfare Association', font_size=11, bold=False)
# Merge right cell visually
cells[1].text = ''

# ══════════════════════════════════════════════
# FOOTER IMAGE
# ══════════════════════════════════════════════
if os.path.exists(footer_img):
    add_para('', font_size=12, space_after=6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(footer_img, width=Inches(6.0))

# ══════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════
output = '/home/jayant/Desktop/Hermes-AgentOS-Dashboard/tabs/mezzaria-style/EC_Minutes_13Jun2026.docx'
doc.save(output)
print(f'Saved: {output}')
