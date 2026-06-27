"""Content file browser API - separate module to avoid syntax conflicts."""
import base64
import io
import os

CONTENT_ROOT = os.path.expanduser("~")
CONTENT_MAX_FILE = 5 * 1024 * 1024  # 5MB max read

CONTENT_IMAGE_EXTS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.svg', '.bmp', '.ico'}
CONTENT_HIDE_DIRS = {'.git', '__pycache__', 'node_modules', '.cache', '.local', '.gnupg', '.ssh'}
CONTENT_HIDE_FILES = {'.DS_Store', '.env', '.env.local', 'id_rsa', 'id_ed25519'}

# Document extraction — return text content for rich formats
CONTENT_DOC_EXTS = {'.pdf', '.docx', '.xlsx', '.xls', '.csv', '.rtf', '.odt', '.ods'}


def _extract_pdf(file_path):
    """Extract text from PDF via PyMuPDF (fitz)."""
    try:
        import fitz
    except ImportError:
        return None
    try:
        doc = fitz.open(file_path)
        pages_text = []
        for page in doc:
            txt = page.get_text("text")
            if txt.strip():
                pages_text.append(txt)
        doc.close()
        if pages_text:
            return "\n\n".join(pages_text)
        # Fallback: OCR-like scan for images
        return f"[PDF: {len(doc)} page(s) — no extractable text found. Try selecting text in the viewer.]"
    except Exception as e:
        return f"[PDF extraction error: {e}]"


def _extract_docx(file_path):
    """Extract text from DOCX via python-docx."""
    try:
        import docx as docx_mod
    except ImportError:
        return None
    try:
        doc = docx_mod.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(c for c in cells if c)
                if row_text:
                    paragraphs.append(row_text)
        if paragraphs:
            return "\n".join(paragraphs)
        return "[DOCX: Document is empty]"
    except Exception as e:
        return f"[DOCX extraction error: {e}]"


def _extract_xlsx(file_path):
    """Extract text from XLSX/ODS via openpyxl."""
    try:
        import openpyxl
    except ImportError:
        return None
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        sheets_text = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                line = " | ".join(c for c in cells if c.strip())
                if line.strip():
                    rows_text.append(line)
            if rows_text:
                sheets_text.append(f"=== {sheet_name} ===\n" + "\n".join(rows_text))
        wb.close()
        if sheets_text:
            return "\n\n".join(sheets_text)
        return "[XLSX: Spreadsheet is empty]"
    except Exception as e:
        return f"[XLSX extraction error: {e}]"


def content_safe_path(requested):
    """Resolve path and ensure it stays within CONTENT_ROOT."""
    if not requested:
        return CONTENT_ROOT
    requested = os.path.expanduser(requested)
    if not requested.startswith('/'):
        requested = os.path.join(CONTENT_ROOT, requested)
    resolved = os.path.realpath(requested)
    if not resolved.startswith(os.path.realpath(CONTENT_ROOT)):
        return None
    return resolved


def content_list_directory(dir_path):
    """List directory entries with metadata."""
    entries = []
    try:
        for name in sorted(os.listdir(dir_path)):
            if name.startswith('.') and name in CONTENT_HIDE_DIRS:
                continue
            if name in CONTENT_HIDE_FILES:
                continue
            full_path = os.path.join(dir_path, name)
            try:
                st = os.stat(full_path)
            except OSError:
                continue
            is_dir = os.path.isdir(full_path)
            ext = '' if is_dir else os.path.splitext(name)[1].lower()
            entries.append({
                "name": name,
                "path": full_path,
                "is_dir": is_dir,
                "size": st.st_size if not is_dir else 0,
                "ext": ext,
                "modified": st.st_mtime,
            })
    except PermissionError:
        return None
    return entries


def content_get_file(file_path):
    """Read a file and return its content + type info."""
    try:
        st = os.stat(file_path)
    except OSError as e:
        return {"error": str(e)}

    if st.st_size > CONTENT_MAX_FILE:
        return {"type": "binary", "size": st.st_size, "ext": os.path.splitext(file_path)[1].lower()}

    ext = os.path.splitext(file_path)[1].lower()

    # Document extraction route
    if ext in CONTENT_DOC_EXTS:
        extracted = None
        if ext == '.pdf':
            extracted = _extract_pdf(file_path)
        elif ext == '.docx':
            extracted = _extract_docx(file_path)
        elif ext in ('.xlsx', '.xls'):
            extracted = _extract_xlsx(file_path)
        if extracted is not None:
            return {
                "type": "document",
                "content": extracted,
                "ext": ext,
                "path": file_path,
                "size": st.st_size,
            }
        # Fallback to text if extraction failed
        return {"type": "text", "content": f"[Document preview not available — extraction library for {ext} not loaded]", "ext": ext, "path": file_path, "size": st.st_size}

    if ext in CONTENT_IMAGE_EXTS:
        try:
            with open(file_path, 'rb') as f:
                data = f.read(st.st_size)
            mime_map = {
                '.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                '.gif': 'image/gif', '.webp': 'image/webp', '.svg': 'image/svg+xml',
            }
            return {
                "type": "image",
                "mime": mime_map.get(ext, 'application/octet-stream'),
                "content": base64.b64encode(data).decode(),
                "ext": ext,
                "path": file_path,
            }
        except Exception as e:
            return {"error": str(e)}

    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
    except Exception as e:
        return {"error": str(e)}

    # Detect HTML files for iframe preview
    if ext in ('.html', '.htm', '.svg'):
        return {
            "type": "html",
            "content": content,
            "ext": ext,
            "path": file_path,
            "size": st.st_size,
        }

    return {
        "type": "text",
        "content": content,
        "ext": ext,
        "path": file_path,
        "size": st.st_size,
    }


def content_save_file(file_path, content):
    """Save content to a file."""
    try:
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return {"ok": True}
    except Exception as e:
        return {"ok": False, "error": str(e)}
